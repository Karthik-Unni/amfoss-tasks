from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ConversationHandler, CallbackQueryHandler
import csv
import requests
import os
from docx import Document

TOKEN = ""
APIKEY = ""
GENRE, BOOK_TITLE, BOOK_NAME, CHOOSING_ACTION = range(4)
def fetch_books_by_genre(genre):
    api_url = f"https://www.googleapis.com/books/v1/volumes?q=subject:{genre}&key={APIKEY}&maxResults=10"
    response = requests.get(api_url)
    books = response.json().get("items", [])
    return books

async def start_command(update: Update, context):
    await update.message.reply_text(
        "Welcome to PagePal! The All in One Telegram Book Store\n\n"
        "Use these commands:\n"
        "- /book to get book by genre.\n"
        "- /preview to get a preview of a book.\n"
        "- /list to add or delete books to your reading list.\n"
        "- /help to get all the commands available.\n"
        "Have a blast!!"
    )

async def book_command(update: Update, context):
    await update.message.reply_text("Please type the genre of books you want to see:")
    return GENRE

async def genre_handler(update: Update, context):
    genre = update.message.text
    books = fetch_books_by_genre(genre)

    with open("books.csv", "w", newline='') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(["Title", "Author", "Description", "Year Published", "Language", "Preview Link"])
        for book in books:
            info = book["volumeInfo"]
            writer.writerow([
                info.get("title", "Unknown"),
                ", ".join(info.get("authors", ["Unknown"])),
                info.get("description", "No description available"),
                info.get("publishedDate", "Unknown"),
                info.get("language", "Unknown"),
                info.get("previewLink", "No preview available")
            ])

    await update.message.reply_document(open("books.csv", "rb"))
    return ConversationHandler.END

async def preview_command(update: Update, context):
    await update.message.reply_text("Please type the title of the book to get a preview link:")
    return BOOK_TITLE

async def book_title_handler(update: Update, context):
    title = update.message.text
    api_url = f"https://www.googleapis.com/books/v1/volumes?q=intitle:{title}&key={APIKEY}&maxResults=1"
    response = requests.get(api_url)
    books = response.json().get("items", [])

    if books:
        preview_link = books[0]["volumeInfo"].get("previewLink", "No preview available")
        await update.message.reply_text(f"Preview link: {preview_link}")
    else:
        await update.message.reply_text("No preview available for this book.")
    
    return ConversationHandler.END

async def list_command(update: Update, context):
    await update.message.reply_text("Please type the name of the book to manage in your reading list:")
    return BOOK_NAME

async def reading_list_command(update: Update, context):
    keyboard = [
        [InlineKeyboardButton("Add a book", callback_data='add')],
        [InlineKeyboardButton("Delete a book", callback_data='delete')],
        [InlineKeyboardButton("View Reading List", callback_data='view')]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text("Choose an action:", reply_markup=reply_markup)
    return CHOOSING_ACTION

async def button_handler(update: Update, context):
    query = update.callback_query
    await query.answer()
    
    if query.data == 'add':
        await query.edit_message_text(text="Type the name of the book to add:")
        return BOOK_NAME
    elif query.data == 'delete':
        await query.edit_message_text(text="Type the name of the book to delete:")
        return BOOK_NAME
    elif query.data == 'view':
        await view_reading_list(update, context)

async def add_book_to_list(book_name, preview_link=None):
    doc = Document()
    if os.path.exists("reading_list.docx"):
        doc = Document("reading_list.docx")
    doc.add_paragraph(f"Title: {book_name}")
    if preview_link:
        doc.add_paragraph(f"Preview: {preview_link}")
    doc.save("reading_list.docx")

async def view_reading_list(update: Update, context):
    if os.path.exists("reading_list.docx"):
        await update.effective_message.reply_document(open("reading_list.docx", "rb"))
    else:
        await update.effective_message.reply_text("Your reading list is empty.")

async def add_book_handler(update: Update, context):
    book_name = update.message.text
    await add_book_to_list(book_name)
    await update.message.reply_text(f"{book_name} has been added to your reading list.")
    return ConversationHandler.END

async def delete_book_handler(update: Update, context):
    book_name = update.message.text
    doc = Document()
    if os.path.exists("reading_list.docx"):
        doc = Document("reading_list.docx")
        for para in doc.paragraphs:
            if book_name in para.text:
                doc.paragraphs.remove(para)
                doc.save("reading_list.docx")
                await update.message.reply_text(f"{book_name} has been removed from your reading list.")
                return ConversationHandler.END
    await update.message.reply_text(f"{book_name} not found in your reading list.")
    return ConversationHandler.END

async def help_command(update: Update, context):
    await update.message.reply_text(
        "Available commands:\n"
        "- /start: Start the bot and see options.\n"
        "- /book: Get book by genre.\n"
        "- /preview: Get a preview of a book.\n"
        "- /list: Add or delete books to your reading list.\n"
        "- /reading_list: View your reading list.\n"
        "- /help: View all the commands."
    )

def main():
    application = Application.builder().token(TOKEN).build()
    book_conv_handler = ConversationHandler(
        entry_points=[CommandHandler("book", book_command)],
        states={GENRE: [MessageHandler(filters.TEXT, genre_handler)]},
        fallbacks=[]
    )
    
    preview_conv_handler = ConversationHandler(
        entry_points=[CommandHandler("preview", preview_command)],
        states={BOOK_TITLE: [MessageHandler(filters.TEXT, book_title_handler)]},
        fallbacks=[]
    )

    list_conv_handler = ConversationHandler(
        entry_points=[CommandHandler("list", list_command)],
        states={
            BOOK_NAME: [
                MessageHandler(filters.TEXT, add_book_handler),
                MessageHandler(filters.TEXT, delete_book_handler)
            ],
        },
        fallbacks=[]
    )
    
    reading_list_conv_handler = ConversationHandler(
        entry_points=[CommandHandler("reading_list", reading_list_command)],
        states={
            CHOOSING_ACTION: [CallbackQueryHandler(button_handler)],
            BOOK_NAME: [
                MessageHandler(filters.TEXT, add_book_handler),
                MessageHandler(filters.TEXT, delete_book_handler)
            ]
        },
        fallbacks=[]
    )

    application.add_handler(CommandHandler("start", start_command))
    application.add_handler(CommandHandler("help", help_command))
    application.add_handler(book_conv_handler)
    application.add_handler(preview_conv_handler)
    application.add_handler(list_conv_handler)
    application.add_handler(reading_list_conv_handler)
    application.run_polling()

if __name__ == "__main__":
    main()
